from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_progress_report():
    doc = Document()

    # --- Title Section ---
    title = doc.add_heading('ACCESS Progress Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Project Information
    p = doc.add_paragraph()
    p.add_run('Project Title: ').bold = True
    p.add_run('A Multimodal Transformer Framework for 3D Scene Understanding in Natural Disaster Assessment')
    
    p = doc.add_paragraph()
    p.add_run('Principal Investigator: ').bold = True
    p.add_run('Nhut Le (Student Lead), Prof. Maryam Rahnemoonfar (Advisor)')
    
    p = doc.add_paragraph()
    p.add_run('Supporting Grant: ').bold = True
    p.add_run('NSF #2423211 "PFI (MCA): Smart Disaster Response..."')

    doc.add_heading('1. Summary of Work Accomplished', level=1)
    
    intro = doc.add_paragraph('During this reporting period, we successfully utilized ACCESS resources to advance the state-of-the-art in post-disaster scene understanding. Our work has focused on three key pillars: the creation of a novel 3D benchmark dataset, the development of a learnable ordering mechanism for 3D point transformers, and a geometric approach to flood depth estimation.')

    # Subsection A
    doc.add_heading('A. Development of 3DAeroRelief: A 3D Benchmark for Post-Disaster Assessment', level=2)
    p = doc.add_paragraph('To address the scarcity of 3D data in disaster scenarios, we constructed 3DAeroRelief, the first 3D benchmark dataset specifically tailored for post-disaster assessment.')
    
    # List for A
    items_a = [
        "Data Construction: We utilized Structure-from-Motion (SfM) and Multi-View Stereo (MVS) techniques to reconstruct dense 3D point clouds from UAV footage captured in the wake of Hurricane Ian (2022).",
        "Scale and Complexity: The dataset features high-fidelity point clouds (avg. 0.52 million points per scan) covering complex outdoor environments, explicitly distinguishing between damaged and intact infrastructure.",
        "Benchmarking: We benchmarked multiple state-of-the-art models (e.g., PTv3, OA-CNNs) on this dataset, utilizing ACCESS GPU resources to establish baseline performance metrics for the research community."
    ]
    for item in items_a:
        doc.add_paragraph(item, style='List Bullet')

    # Subsection B
    doc.add_heading('B. OPTNet: Ordering Point Transformer Network', level=2)
    p = doc.add_paragraph('We developed OPTNet, a novel architecture designed to overcome the limitations of static serialization (e.g., Hilbert curves) in processing irregular post-disaster 3D point clouds.')
    
    # List for B
    items_b = [
        "Innovation: We introduced a \"Point Sorter\" module that dynamically predicts an optimal permutation for point cloud serialization, maximizing the locality of attention mechanisms.",
        "Methodology: The model was trained using a self-supervised ordering loss, optimizing for both locality (grouping spatial neighbors) and uniformity (preventing mode collapse).",
        "Results: OPTNet achieved a state-of-the-art mIoU of 79.65% on the 3DAeroRelief dataset, significantly outperforming the PTv3 baseline (45.84%), with particular success in segmenting elongated structures like roads."
    ]
    for item in items_b:
        doc.add_paragraph(item, style='List Bullet')

    # Subsection C
    doc.add_heading('C. Geometric Flood Depth Estimation', level=2)
    p = doc.add_paragraph('We proposed and validated a pipeline for estimating flood depth from monocular aerial imagery by fusing 2D segmentation with Digital Elevation Models (DEMs).')
    
    # List for C
    items_c = [
        "Approach: We utilized Mask2Former, fine-tuned on ACCESS resources (Jetstream2), to generate precise 2D flood masks. These masks were intersected with DEMs to calculate the \"Water Surface Elevation\" based on the principle of local hydrostatic equilibrium.",
        "Validation: The method was validated on the FloodNet and CRASAR-U-DROIDS datasets, demonstrating the ability to extract volumetric data without the high latency of hydrodynamic simulations."
    ]
    for item in items_c:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2. Resource Utilization', level=1)
    p = doc.add_paragraph('ACCESS computational resources were critical to the success of these computationally intensive tasks:')
    
    resources = [
        "High-Performance Training: The training of the OPTNet architecture required significant compute power. We utilized NVIDIA H100 GPUs to train the models for 3000 epochs, enabling the optimization of the learnable Point Sorter module.",
        "Large-Scale Benchmarking: The evaluation of baseline models (PTv1, PTv2, PTv3, FPT, OA-CNNs) for the 3DAeroRelief dataset paper was conducted on 8 NVIDIA A5000 GPUs, ensuring consistent high-performance computation for processing dense point clouds.",
        "Model Fine-Tuning: For the flood depth estimation project, we utilized Jetstream2 at Indiana University (Allocation CIS251047) to fine-tune the Mask2Former architecture on the FloodNet dataset at 1024x1024 resolution.",
        "Data Processing: The reconstruction of high-resolution 3D point clouds from UAV footage and the storage of the resulting datasets required the utilization of the allocated storage resources."
    ]
    for item in resources:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. Publications and Products', level=1)
    p = doc.add_paragraph('The ACCESS-supported work has resulted in three manuscripts currently under submission/review:')

    # Publication 1
    p = doc.add_paragraph()
    p.add_run('1. OPTNet: Ordering Point Transformer Network for Post-disaster 3D Semantic Segmentation').bold = True
    doc.add_paragraph('Authors: Nhut Le, Ehsan Karimi, Maryam Rahnemoonfar')
    doc.add_paragraph('Status: Submitted to the International Conference on Pattern Recognition (ICPR) 2026.')

    # Publication 2
    p = doc.add_paragraph()
    p.add_run('2. 3DAeroRelief: The first 3D Benchmark UAV Dataset for Post-Disaster Assessment').bold = True
    doc.add_paragraph('Authors: Nhut Le, Ehsan Karimi, Maryam Rahnemoonfar')
    doc.add_paragraph('Status: Submitted to Scientific Data.')

    # Publication 3
    p = doc.add_paragraph()
    p.add_run('3. Geometric Flood Depth Estimation: Fusing Transformer-Based Segmentation with Digital Elevation Models').bold = True
    doc.add_paragraph('Authors: Nhut Le, Ehsan Karimi, Maryam Rahnemoonfar')
    doc.add_paragraph('Status: Submitted to IEEE International Geoscience and Remote Sensing Symposium (IGARSS) 2026.')

    doc.add_heading('4. Plan for Remaining Resources', level=1)
    p = doc.add_paragraph('We plan to utilize the remaining ACCESS credits to complete the objectives outlined in our original proposal:')
    
    plans = [
        "Multimodal Fusion: We will extend OPTNet to incorporate the multimodal fusion of RGB imagery and LiDAR point clouds, as described in the \"Multimodal Transformer Framework\" section of our proposal. This will require extensive GPU hours to train the cross-modal attention mechanisms.",
        "Damage Assessment Integration: Building on the qualitative results from the flood depth estimation project, we will integrate volumetric depth data into damage assessment networks to quantify structural risk for individual buildings.",
        "Expanded Benchmarking: We will continue to evaluate and refine our models on the 3DAeroRelief dataset to ensure robustness across different geographic zones and disaster types."
    ]
    for item in plans:
        doc.add_paragraph(item, style='List Bullet')

    # Save the document
    file_name = 'ACCESS_Progress_Report.docx'
    doc.save(file_name)
    print(f"Document saved as {file_name}")

if __name__ == "__main__":
    create_progress_report()